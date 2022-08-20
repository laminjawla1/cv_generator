from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

speak("Hello and welcome to Sage's CV generator, version 0.1")

# Importing the profile picture
document.add_picture("sage.jpg", width=Inches(2.0))

# Taking details from the user
name = input("Name: ")
phone_number = input("Phone Number: ")
email = input("Email: ")

# Writing to the file
document.add_paragraph(name + " | " + phone_number + " | " + email)

# About me
document.add_heading("About me")
document.add_paragraph(input("Tell me about yourself: "))

# Work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Company name: ")
from_date = input("From date: ")
to_date = input("To date: ")

p.add_run(company + " ").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True

experience_details = input("Describe your experience at :" + company)
p.add_run(experience_details)

# More experiences
while True:
    has_more_experience = input("Do you have more experiences? (yes or no): ")
    if has_more_experience.lower() == "yes":
        p = document.add_paragraph()

        company = input("Company name: ")
        from_date = input("From date: ")
        to_date = input("To date: ")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True

        experience_details = input("Describe your experience at " + company + ": ")
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading("SKILLS")
skill = input("Sill: ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

while True:
    check = input("Do you have more skills? Yes or No: ")
    if check.lower() == "yes":
        skill = input("Sill: ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using -Sage CV generator version 0.1"

# Saving the file
document.save("my_cv.docx")
print("Info Save!")
quit()
